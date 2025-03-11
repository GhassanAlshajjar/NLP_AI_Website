import docx
import pdfplumber
import re

def clean_text(text):
    """Normalizes extracted text: lowercase, remove extra spaces, and special characters."""
    text = text.lower()  # Convert to lowercase
    text = re.sub(r'\s+', ' ', text).strip()  # Remove extra spaces and newlines
    return text

def extract_text(file):
    """
    Extracts text from PDF, DOCX, or TXT files and normalizes it.
    """
    try:
        text = ""

        if file.filename.endswith('.pdf'):
            with pdfplumber.open(file) as pdf:
                text = "\n".join([page.extract_text() or "" for page in pdf.pages])

        elif file.filename.endswith('.docx'):
            doc = docx.Document(file)
            extracted_text = []

            for para in doc.paragraphs:
                if para.text.strip():
                    extracted_text.append(para.text)

            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        extracted_text.append(row_text)

            for section in doc.sections:
                if section.header:
                    header_text = section.header.paragraphs[0].text.strip()
                    if header_text:
                        extracted_text.append(f"Header: {header_text}")

                if section.footer:
                    footer_text = section.footer.paragraphs[0].text.strip()
                    if footer_text:
                        extracted_text.append(f"Footer: {footer_text}")

            text = "\n".join(extracted_text)

        elif file.filename.endswith('.txt'):
            text = file.read().decode('utf-8')

        return clean_text(text)  # Normalize text before returning

    except Exception as e:
        print(f"Error extracting text: {e}")
        return ""
